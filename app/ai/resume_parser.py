from app.ai import openai_client
from app.ai.prompt_templates import RESUME_PARSE_PROMPT
from app.schemas.resume import ParsedResume
from app.utils.logger import get_logger
import io

logger = get_logger(__name__)


async def extract_text_from_pdf(file_bytes: bytes) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text_parts.append(extracted)
    return "\n".join(text_parts)


async def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


async def parse_resume(file_bytes: bytes, file_extension: str) -> tuple[str, ParsedResume]:
    if file_extension == "pdf":
        raw_text = await extract_text_from_pdf(file_bytes)
    elif file_extension in ("docx", "doc"):
        raw_text = await extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

    if not raw_text.strip():
        raise ValueError("Could not extract text from the uploaded file")

    logger.info("parsing_resume", text_length=len(raw_text))

    prompt = RESUME_PARSE_PROMPT.format(resume_text=raw_text[:8000])
    messages = [
        {"role": "system", "content": "You are an expert resume parser. Always respond with valid JSON."},
        {"role": "user", "content": prompt},
    ]

    parsed_data = await openai_client.chat_completion_json(messages, temperature=0.2)

    resume = ParsedResume(
        skills=parsed_data.get("skills", []),
        projects=parsed_data.get("projects", []),
        experience=parsed_data.get("experience", []),
        education=parsed_data.get("education", []),
        certifications=parsed_data.get("certifications", []),
        research_papers=parsed_data.get("research_papers", []),
        achievements=parsed_data.get("achievements", []),
        summary=parsed_data.get("summary"),
    )

    logger.info(
        "resume_parsed",
        skills_count=len(resume.skills),
        projects_count=len(resume.projects),
        experience_count=len(resume.experience),
    )

    return raw_text, resume
